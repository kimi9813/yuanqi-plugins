"""
生成《yuanqi-plugins 部署方法》Word 文档
运行: python scripts/generate_deployment_doc.py
输出: docs/部署方法.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/部署方法.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)


def set_chinese_font(run, font_name="Microsoft YaHei", size=10.5, bold=False, color=None):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_chinese_font(run, size=(18 if level == 1 else (14 if level == 2 else 12)), bold=True)
    return p


def add_paragraph(doc, text, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def add_code_block(doc, code, language="bash"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(code)
    set_chinese_font(run, font_name="Consolas", size=9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("说明：" + text)
    set_chinese_font(run, size=9.5, color=(0xE6, 0x74, 0x00))
    return p


def main():
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("yuanqi-plugins 部署方法")
    set_chinese_font(run, size=22, bold=True, color=(0x1A, 0x25, 0x7A))

    add_paragraph(doc, "本文档介绍 yuanqi-plugins 项目的多种公网部署方式，以及部署后注册到腾讯元器插件中心的完整流程。", size=10.5)

    # 目录/前置
    add_heading(doc, "一、前置准备", level=1)
    add_paragraph(doc, "1. 代码已推送到 GitHub 仓库：https://github.com/kimi9813/yuanqi-plugins")
    add_paragraph(doc, "2. 本地已验证 `python app.py` 可正常启动，访问 `http://localhost:8000/docs` 可查看 API 文档。")
    add_paragraph(doc, "3. 准备一个可公网访问的服务器或 PaaS 平台账号（Render / Railway / Fly.io / GitHub Codespaces 等）。")
    add_paragraph(doc, "4. 部署完成后，需要将 `specs/*.yaml` 中的 `servers.url` 替换为实际公网地址，再上传到腾讯元器插件中心。")

    # 通用流程
    add_heading(doc, "二、通用部署后配置", level=1)
    add_paragraph(doc, "无论使用哪种部署方式，完成后都需要执行以下两步：")
    add_paragraph(doc, "步骤 1：获取公网地址。例如：https://your-app.example.com")
    add_paragraph(doc, "步骤 2：修改 5 个 OpenAPI 规范文件中的 servers.url。")
    add_code_block(doc, """# 将 specs/web_tool_openapi.yaml 中的
servers:
  - url: http://localhost:8000/web
# 替换为
servers:
  - url: https://your-app.example.com/web
""")
    add_paragraph(doc, "步骤 3：在腾讯元器插件中心，逐个上传 `specs/*.yaml` 文件完成插件注册。")

    # 方法1：GitHub Codespaces
    add_heading(doc, "三、方法 1：GitHub Codespaces（免费，适合测试）", level=1)
    add_paragraph(doc, "优点：零配置、免费、与 GitHub 仓库集成。缺点：Codespace 停止后 URL 可能变化，不适合生产长期运行。", size=10)
    steps = [
        "打开 GitHub 仓库页面，点击 Code → Codespaces → Create codespace on master。",
        "等待环境初始化完成，终端会自动进入 `/workspaces/yuanqi-plugins`。",
        "运行以下命令安装依赖并启动服务：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, "pip install -r requirements.txt\npython app.py")
    add_paragraph(doc, "等待出现 `Uvicorn running on http://0.0.0.0:8000` 后，点击右下角「转发端口 8000」。")
    add_paragraph(doc, "右键 forwarded port，选择「端口可见性」→「Public」，复制生成的公网 URL（形如 `https://xxx-8000.app.github.dev`）。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 中的 servers.url，注意去掉末尾多余斜杠，保持 `/web`、 `/file` 等前缀。")

    # 方法2：Render
    add_heading(doc, "四、方法 2：Render（免费 Web Service，推荐）", level=1)
    add_paragraph(doc, "优点：免费套餐支持 24/7 运行，配置简单。缺点：免费实例会在 15 分钟无访问后休眠，首次访问有冷启动延迟。", size=10)
    steps = [
        "访问 https://dashboard.render.com/ 并注册/登录。",
        "点击 New → Web Service，选择 GitHub 仓库 `kimi9813/yuanqi-plugins`。",
        "Render 会自动识别根目录的 `render.yaml`，填写以下信息：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, """Name: yuanqi-plugins
Region: Singapore (Asia Pacific)
Branch: master
Runtime: Python 3
Build Command: pip install -r requirements.txt
Start Command: uvicorn app:app --host 0.0.0.0 --port 10000
""")
    add_paragraph(doc, "点击 Create Web Service，等待部署完成。")
    add_paragraph(doc, "部署成功后，Render 会分配一个 `https://yuanqi-plugins-xxx.onrender.com` 的域名。")
    add_paragraph(doc, "在 Dashboard → Disks 中创建 Disk，Name 填 `data`，Mount Path 填 `/app/data`，Size 填 `1 GB`，并挂载到该服务。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册到腾讯元器。")

    # 方法3：Railway
    add_heading(doc, "五、方法 3：Railway（按量计费，启动快）", level=1)
    add_paragraph(doc, "优点：部署简单、启动速度快、支持 Docker。缺点：免费额度有限，长期使用需付费。", size=10)
    steps = [
        "访问 https://railway.app/ 并登录。",
        "点击 New Project → Deploy from GitHub repo，选择 `yuanqi-plugins`。",
        "Railway 会自动读取 `railway.json` 中的配置。",
        "在 Variables 中添加环境变量（如需要）：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, "PORT=8000")
    add_paragraph(doc, "等待部署完成，Railway 会生成 `https://yuanqi-plugins.up.railway.app` 的域名。")
    add_paragraph(doc, "在 Settings → Networking 中可绑定自定义域名。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册。")

    # 方法4：Fly.io
    add_heading(doc, "六、方法 4：Fly.io（按量计费，全球边缘）", level=1)
    add_paragraph(doc, "优点：全球边缘节点、容器化部署、按量付费。缺点：需要信用卡验证。", size=10)
    steps = [
        "安装 Fly CLI：https://fly.io/docs/hands-on/install-flyctl/",
        "登录账号：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, "flyctl auth login")
    add_paragraph(doc, "在项目目录下初始化并部署：")
    add_code_block(doc, """cd yuanqi-plugins
flyctl launch --name yuanqi-plugins --region hkg
flyctl deploy
""")
    add_paragraph(doc, "首次部署会创建 Volume，按提示执行：")
    add_code_block(doc, "flyctl volumes create yuanqi_data --size 1 --region hkg")
    add_paragraph(doc, "部署完成后，Fly.io 会分配 `https://yuanqi-plugins.fly.dev`。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册。")

    # 方法5：VPS/Docker
    add_heading(doc, "七、方法 5：VPS + Docker（最稳定，适合生产）", level=1)
    add_paragraph(doc, "优点：完全可控、性能稳定。缺点：需要一台服务器和基础运维能力。", size=10)
    steps = [
        "准备一台 Linux 服务器（Ubuntu 22.04 推荐），安装 Docker 和 Docker Compose。",
        "克隆代码到服务器：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, "git clone https://github.com/kimi9813/yuanqi-plugins.git\ncd yuanqi-plugins")
    add_paragraph(doc, "使用 Docker Compose 启动：")
    add_code_block(doc, "docker compose up -d --build")
    add_paragraph(doc, "服务将监听服务器的 8000 端口。若需 80/443，可搭配 Nginx 反向代理或 `docker compose` 端口映射。")
    add_paragraph(doc, "更新代码时执行：")
    add_code_block(doc, "git pull && docker compose up -d --build")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册。")

    # 方法6：阿里云 FC
    add_heading(doc, "八、方法 6：阿里云函数计算 FC（函数计算，国内访问快）", level=1)
    add_paragraph(doc, "优点：按调用计费、不用维护服务器、国内访问快。缺点：实例无状态，本地文件/数据在实例回收后丢失。", size=10)
    steps = [
        "访问 https://fc.console.aliyun.com/ 并开通函数计算服务。",
        "点击「函数」→「创建函数」，选择「自定义创建」：",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_code_block(doc, """函数类型：HTTP 函数
函数名称：yuanqi-plugins
运行环境：Python 3.10
请求处理程序（Handler）：handler
内存规格：1024 MB
超时时间：60 秒
""")
    add_paragraph(doc, "上传代码：可下载 GitHub 仓库 ZIP 包，在「代码」标签页上传；或使用 Serverless Devs 命令行 `s deploy`。")
    add_paragraph(doc, "创建 HTTP 触发器，请求方法勾选 GET/POST/PUT/DELETE/OPTIONS，认证方式选择「无需认证」。")
    add_paragraph(doc, "部署完成后复制公网地址，例如 `https://yuanqi-plugins-xxx.cn-hangzhou.fcapp.run`。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册。")

    # 方法7：腾讯云 SCF
    add_heading(doc, "九、方法 7：腾讯云云函数 SCF（函数计算）", level=1)
    add_paragraph(doc, "优点：与腾讯元器同属腾讯云生态，访问延迟低。缺点：函数计算实例无状态，本地数据会丢失。", size=10)
    add_paragraph(doc, "推荐选择 Web 函数，无需单独配置 API 网关；如果没有 Web 函数选项，则选择事件函数 + API 网关触发器。")
    add_paragraph(doc, "方式 A：Web 函数（推荐）")
    add_code_block(doc, """函数名称：yuanqi-plugins
函数类型：Web 函数
运行环境：Python 3.10
执行方法：serverless.scf_handler.app
内存：1024 MB
超时时间：60 秒
""")
    add_paragraph(doc, "上传代码：下载 GitHub 仓库 ZIP 包，在函数「代码」页本地上传。")
    add_paragraph(doc, "创建完成后 SCF 会自动分配访问地址，例如 `https://yuanqi-plugins-xxx.gz.apigw.tencentcs.com/`。")
    add_paragraph(doc, "方式 B：事件函数 + API 网关")
    add_code_block(doc, """函数名称：yuanqi-plugins
函数类型：事件函数
运行环境：Python 3.10
执行方法：serverless.scf_handler.main_handler
内存：1024 MB
超时时间：60 秒
触发器：API 网关
请求方法：ANY
""")
    add_paragraph(doc, "上传代码后，进入「触发管理」→「创建触发器」，选择「API 网关」，服务名填 `yuanqi-plugins-api`，发布环境选择「发布」，请求方法选择「ANY」。")
    add_paragraph(doc, "创建完成后复制访问路径，例如 `https://service-xxx.gz.apigw.tencentcs.com/release/`。")
    add_paragraph(doc, "按「二、通用部署后配置」修改 `specs/*.yaml` 并注册。")

    # 注册到腾讯元器
    add_heading(doc, "十、注册到腾讯元器插件中心", level=1)
    steps = [
        "登录腾讯元器后台（https://yuanqi.tencent.com/）。",
        "进入「插件中心」→「创建插件」。",
        "选择「OpenAPI 插件」，上传对应 YAML 文件。",
        "元器会解析 `servers.url` 和 `paths` 下的所有 operationId，每个 operation 对应一个工具。",
        "按提示配置插件名称、描述、图标，保存并发布。",
    ]
    for s in steps:
        add_paragraph(doc, s)
    add_note(doc, "上传前务必确认 servers.url 是 HTTPS 公网地址，且服务处于可访问状态，否则元器校验会失败。")

    # 安全建议
    add_heading(doc, "十一、安全与运维建议", level=1)
    tips = [
        "terminal_tool 支持执行代码和 Shell，建议仅在可信环境或启用访问控制后使用。",
        "生产环境建议增加 API Key 鉴权或仅允许腾讯元器 IP 段访问。",
        "定期清理 data/files 目录，避免磁盘占满。",
        "如果不需要 Java 执行，可从 Dockerfile 中移除 default-jdk 以减小镜像体积。",
        "建议为服务配置 HTTPS（Render/Railway/Fly.io 默认提供，VPS 需自行配置证书）。",
        "函数计算场景下，文件/技能/任务等本地数据会丢失，生产环境请挂载 NAS/CFS 或改用对象存储/数据库。",
        "函数计算建议设置超时 60 秒、内存 1024 MB，避免文件转换和网页抓取超时。",
    ]
    for i, tip in enumerate(tips, 1):
        add_paragraph(doc, f"{i}. {tip}")

    # 常见问题
    add_heading(doc, "十二、常见问题", level=1)
    faq = [
        ("Q：元器解析 YAML 时提示 servers.url 错误？", "A：确保地址以 https:// 开头，且末尾不要重复路径前缀。例如部署域名是 https://abc.com，则 web_tool 的 servers.url 应为 https://abc.com/web。"),
        ("Q：部署后访问 /docs 正常，但元器调用返回 404？", "A：检查 YAML 中的 operationId 是否与代码中的路由一致，且路径前缀正确。"),
        ("Q：文件转换功能报错？", "A：确认服务器安装了相关依赖。Docker 镜像已内置 python-docx、openpyxl、python-pptx、PyPDF2。"),
        ("Q：Java 代码无法执行？", "A：确认服务器已安装 JDK。Docker 镜像已内置 default-jdk。"),
        ("Q：函数计算部署后文件/任务数据丢失？", "A：函数计算实例无状态，本地数据会随实例回收清空。测试场景不影响，生产场景请挂载 NAS/CFS 或改用数据库存储。"),
        ("Q：函数计算返回 502 错误？", "A：通常因超时或内存不足导致。请将超时时间设为 60 秒，内存设为 1024 MB。"),
    ]
    for q, a in faq:
        add_paragraph(doc, q, bold=True)
        add_paragraph(doc, a)

    doc.save(OUTPUT)
    print(f"已生成文档：{OUTPUT.resolve()}")


if __name__ == "__main__":
    main()
